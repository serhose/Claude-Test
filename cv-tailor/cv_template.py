"""
CV Word Template — Melda Akan
Generates a .docx file matching the source CV format using python-docx.
"""

import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_bottom_border(paragraph):
    """Add a bottom border line under a paragraph (section header rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_spacing(paragraph, before=0, after=0):
    pPr = paragraph._p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(before))
    pSpacing.set(qn('w:after'), str(after))
    pPr.append(pSpacing)


def _set_right_tab(paragraph, twips=10800):
    """Set a right-aligned tab stop (default 6.5 inches = 9360 twips)."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(twips))
    tabs.append(tab)
    pPr.append(tabs)


def render_cv(cv_data: dict) -> bytes:
    doc = Document()

    # Page margins (0.5 inch all sides)
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    # Remove default paragraph spacing from Normal style
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)

    personal = cv_data["personal"]

    # ── Header ────────────────────────────────────────────────────────────────

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=0)
    r = p.add_run(personal["name"])
    r.bold = True
    r.font.size = Pt(17)
    r.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=0)
    r = p.add_run(f"Address: {personal['location']}")
    r.font.size = Pt(10)
    r.font.name = 'Calibri'

    contact_str = (
        f"Tel: {personal['phone']}  |  "
        f"Email: {personal['email']}  |  "
        f"{personal['linkedin']}"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=60)
    r = p.add_run(contact_str)
    r.font.size = Pt(10)
    r.font.name = 'Calibri'

    # ── Helpers ───────────────────────────────────────────────────────────────

    def section_header(title):
        p = doc.add_paragraph()
        _set_spacing(p, before=80, after=20)
        _add_bottom_border(p)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Calibri'

    def company_line(left, right):
        p = doc.add_paragraph()
        _set_spacing(p, before=40, after=0)
        _set_right_tab(p)
        r = p.add_run(left)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        p.add_run('\t')
        r = p.add_run(right)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

    def title_line(title, date_str):
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=0)
        _set_right_tab(p)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        p.add_run('\t')
        r = p.add_run(date_str)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

    def bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        _set_spacing(p, before=0, after=0)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

    # ── Experience ────────────────────────────────────────────────────────────

    if cv_data.get("experience"):
        section_header("EXPERIENCE")
        for exp in cv_data["experience"]:
            if "roles" in exp:
                company_line(exp["company"], exp.get("location", ""))
                for role in exp["roles"]:
                    title_line(role["title"], f"{role['start']} \u2013 {role['end']}")
            else:
                company_line(exp["company"], exp.get("location", ""))
                title_line(exp["title"], f"{exp['start']} \u2013 {exp['end']}")
            for b in exp.get("bullets", []):
                bullet(b["text"] if isinstance(b, dict) else b)

    # ── Education ─────────────────────────────────────────────────────────────

    if cv_data.get("education"):
        section_header("EDUCATION")
        for edu in cv_data["education"]:
            gpa = f" ({edu['gpa']} GPA)" if edu.get("gpa") else ""
            company_line(edu["institution"] + gpa, edu.get("location", ""))
            title_line(edu["degree"], f"{edu['start']} \u2013 {edu['end']}")
            for hl in edu.get("highlights", []):
                bullet(hl)
            if edu.get("coursework"):
                bullet("Relevant Coursework: " + ", ".join(edu["coursework"]))

    # ── Volunteer ─────────────────────────────────────────────────────────────

    if cv_data.get("volunteer"):
        section_header("VOLUNTEER WORK")
        for vol in cv_data["volunteer"]:
            company_line(vol["organization"], vol.get("location", ""))
            title_line(vol["title"], f"{vol['start']} \u2013 {vol['end']}")
            for b in vol.get("bullets", []):
                bullet(b["text"] if isinstance(b, dict) else b)

    # ── Skills ────────────────────────────────────────────────────────────────

    if cv_data.get("skills"):
        section_header("SKILLS/ACTIVITIES")
        skills = cv_data["skills"]

        if skills.get("technical"):
            bullet("Technical skills: " + ", ".join(skills["technical"]))
        if skills.get("data"):
            bullet("Data skills: " + "; ".join(skills["data"]))
        if skills.get("finance"):
            bullet("Skills: " + ", ".join(skills["finance"]))
        if skills.get("business"):
            extras = [s for s in skills["business"] if s not in skills.get("finance", [])]
            if extras:
                bullet(", ".join(extras))
        if skills.get("certifications"):
            bullet("Certifications: " + ", ".join(skills["certifications"]))
        if skills.get("languages"):
            lang_str = ", ".join(
                f"{l['language']} ({l['level']})" for l in skills["languages"]
            )
            bullet("Language skills: " + lang_str)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Quick test ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import json, pathlib

    master = json.loads(
        pathlib.Path("master_cv.json").read_text(encoding="utf-8")
    )
    docx_bytes = render_cv(master)
    out = pathlib.Path("test_output.docx")
    out.write_bytes(docx_bytes)
    print(f"Written {len(docx_bytes):,} bytes → {out}")
